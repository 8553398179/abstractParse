# -*- coding: utf-8 -*-
"""
Created on Sun Apr 17 23:09:40 2022

@author: darshanRaghunath
"""

import pandas as pd
df1 = pd.read_csv("C:/Users/darshanRaghunath/bookKalyan1.csv")
col=["8","9"]
df1 = df1[col]




df2 = pd.read_csv("C:/Users/darshanRaghunath/book19.csv")
df2 = df2[col]


 



vocab_size = 5000
embedding_dim = 64
max_length = 100
trunc_type = 'post'
padding_type = 'post'
oov_tok = '<OOV>'
training_portion = .7
num_epochs = 8







#print(df.head())

df1.columns = ['Paragraph', 'class']
df2.columns = ['Paragraph', 'class']
#print(df.head())
#print(df.info())




from sklearn.model_selection import train_test_split
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.naive_bayes import MultinomialNB
from sklearn.metrics import confusion_matrix
from sklearn.metrics import precision_score, recall_score, f1_score, accuracy_score
from sklearn.metrics import precision_recall_fscore_support
import csv
import tensorflow as tf
import numpy as np
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences



#stratify=df['class']



articles = []
labels = []
count=1



noise =0
for i  in df2["Paragraph"]:
        noise += 1
        articles.append(i)    
        if(noise> 3400):
            break
        
        
noise =0    
for i  in df1["Paragraph"]:
        noise += 1
        articles.append(i)    
        if(noise>6900):
            break
        
    
    
    
noise =0        
        
print(len(articles))     
for i in df2["class"]:
    l=[]
    noise += 1
    l.append(i)
    labels.append(l)
    if(noise> 3400):
        break
    
noise =0   
    
for i in df1["class"]:
    l=[]
    noise += 1
    l.append(i)
    labels.append(l)
    if(noise> 6900):
        break
    
    
print(len(labels))



#print(labels[100])
#print(articles[100])

train_articles,validation_articles , train_labels,validation_labels  = train_test_split(articles,labels,test_size=(1-training_portion),random_state = 100)


evalSize = int(len(validation_articles) * 0.3)


evaluate_articles = validation_articles [0:evalSize]
evaluate_labels = validation_labels[0:evalSize]
                                    
validation_articles= validation_articles[evalSize :]
validation_labels =  validation_labels[evalSize :]



print("Validation")
print(len(validation_articles))


print("Evaluation")
print(len(evaluate_articles))

'''

train_size = int(len(articles) * training_portion)
train_articles = articles[0: train_size]
train_labels = labels[0: train_size]

validation_articles = articles[train_size:]
validation_labels = labels[train_size:]

#print(len(validation_labels[0]))

'''



#print(train_size)
print("train and validation")

print(len(train_articles))
print(len(train_labels))
print(len(validation_articles))
print(len(validation_labels))




tokenizer = Tokenizer(num_words = vocab_size, oov_token=oov_tok,filters='+-',)
tokenizer.fit_on_texts(train_articles)
word_index = tokenizer.word_index
#print(word_index)
#dict(list(word_index.items())[0:10])


train_sequences = tokenizer.texts_to_sequences(train_articles)
#print(train_sequences[100])




train_padded = pad_sequences(train_sequences, maxlen=max_length, padding=padding_type, truncating=trunc_type)
print(len(train_sequences[0]))
print(len(train_padded[0]))

print(len(train_sequences[1]))
print(len(train_padded[1]))

print(len(train_sequences[10]))
print(len(train_padded[10]))









validation_sequences = tokenizer.texts_to_sequences(validation_articles)
validation_padded = pad_sequences(validation_sequences, maxlen=max_length, padding=padding_type, truncating=trunc_type)



evaluate_sequences = tokenizer.texts_to_sequences(evaluate_articles)
evaluate_padded = pad_sequences(evaluate_sequences, maxlen=max_length, padding=padding_type, truncating=trunc_type)

print(len(validation_sequences))
print(validation_padded.shape)



reverse_word_index = dict([(value, key) for (key, value) in word_index.items()])

def decode_article(text):
    return ' '.join([reverse_word_index.get(i, '?') for i in text])
print(decode_article(train_padded[100]))
print('---')
print(train_articles[100])

print()
print()


print(decode_article(train_padded[5005]))
print('---')
print(train_articles[5005])



train_labels =np.array(train_labels)
validation_labels = np.array(validation_labels)



evaluate_labels= np.array(evaluate_labels)




'''
#X_train, X_test, y_train, y_test = train_test_split(df['Paragraph'], df['class'],test_size=0.1)
#print(X_train[10:20])
#print(y_train[10:20])


train = df[0:3000]
test = df[3000:]
print(train.head())
print(train.info())

print(test.head())
print(test.info())


#print(train[0])

#print(len(X_train)) 
#print(type(X_train[2]))
import csv
import tensorflow as tf
import numpy as np
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences

'''




model = tf.keras.Sequential([
    # Add an Embedding layer expecting input vocab of size 5000, and output embedding dimension of size 64 we set at the top
    tf.keras.layers.Embedding(vocab_size, embedding_dim),
    tf.keras.layers.Bidirectional(tf.keras.layers.LSTM(embedding_dim)),
    #tf.keras.layers.Bidirectional(tf.keras.layers.LSTM(32)),
    # use ReLU in place of tanh function since they are very good alternatives of each other.
    tf.keras.layers.Dense(embedding_dim, activation='relu'),
    # Add a Dense layer with 6 units and softmax activation.
    # When we have multiple outputs, softmax convert outputs layers into a probability distribution.
    tf.keras.layers.Dense(7, activation='softmax')
])
model.summary()

#print(train_labels.size)
model.compile(loss='sparse_categorical_crossentropy', optimizer='adam', metrics=['accuracy'])


history = model.fit(train_padded, train_labels, epochs=num_epochs, validation_data=(validation_padded, validation_labels), verbose=2)

model.save("E:/model/pvp")

'''
y_prd =model.predict(validation_padded)
print(len(validation_padded))
print(len(y_prd))
print(validation_labels[0])
#y_prd1 =np.argmax(y_prd)
print(np.argmax(y_prd[0]))
#print(confusion_matrix(validation_labels, y_prd1))
'''


y_prd =model.predict(evaluate_padded)
#print(len(validation_padded))
#print(len(y_prd))
#print(validation_labels[0])
#y_prd1 =np.argmax(y_prd)
#print(np.argmax(y_prd[0]))
#print(confusion_matrix(validation_labels, y_prd1))





ycorrect = []
ypred = []

for i in range(len(evaluate_padded)):
    ycorrect.append(evaluate_labels[i][0])
    ypred.append(np.argmax(y_prd[i]))
    
    
print(confusion_matrix(ycorrect, ypred ))



mat = confusion_matrix(ycorrect, ypred )





percent=[]
for i in range(len(mat)):
    denum=0
    num=0
    for j in range(len(mat[i])):
        if(i ==j):
            num = mat[i][i]
        
        denum += mat[i][j]
    percent.append(num/denum)
        
print(percent)


output =["Title", "Author", "Affiliation" , "Abstract", "Noise", "Image" ,"ID's"]
print("length of the training sample",len(train_articles) )
print("length of the testing sample" , len(evaluate_articles))

print()
print()

for i in range(len(percent)):
    print(output[i], percent[i])
    
    
    
import fitz
import docx

import pandas as pd


'''
Change the path file path
change filename in row.append
change the output file name in the last line
'''



#doc1 = fitz.open("C:/Users/darshanRaghunath/pdfOriginal.pdf")

doc1 = fitz.open("C:/Users/darshanRaghunath/Downloads/DocSample.pdf")


Filename = "DocSample"
Final=[]
document =  docx.Document()
pageNumber = 0


text = []
characterists = []


for page in doc1:
        pageNumber += 1 
        #print("H")
        blocks = page.getText("dict")["blocks"]
        for b in blocks:
            paragraph = ""
            character = []
            if b['type'] == 0:  # block contains text
                for l in b["lines"]:  # iterate through the text lines
                    #print("line")
                    #print(l)
                    for s in l["spans"]:  
                        ch = []
                        ch.append(int(s["size"]))
                        ch.append(s["font"])
                        ch.append(s["color"])
                        paragraph = paragraph + (s["text"]) + " "
                        
                        #print(paragraph)
                        #print(character)
                        character.append(ch)
                        
                characterists.append(character)            
                text.append(paragraph)
            
       
                    
                        
    
rules =[]


for i in range(len(characterists)):
    subRule =[]
    for k in range(len(characterists[i])):      
        if(characterists[i][k] not in subRule):
            subRule.append(characterists[i][k])
    rules.append(subRule)
            
            
    '''
        
            
        for j in range(len(characterists[i])):
            if(j != k):
                if(characterists[i][k] == characterists[i][j]):
    '''
                    
 
    


text_1 = tokenizer.texts_to_sequences(text)
text_2 = pad_sequences(text_1, maxlen=max_length, padding=padding_type, truncating=trunc_type)

#print(rules)

TEXT_CLASS =model.predict(text_2)
#print(TEXT_CLASS)
textIndex = []
textLabels = []
for i in range(len(TEXT_CLASS)):
    textIndex.append(np.argmax(TEXT_CLASS[i]))
    textLabels.append(output[np.argmax(TEXT_CLASS[i])])
    
    
    
    
#print(textLabels )
    
abstract = []
title = [ ]
author = []
affliation = []
noise = []
ids = []


for i in range(len(text)):
    #print(text[i])
    #print(textLabels[i])
    print(rules[i])
    if(textIndex [i] == 0):
        title.append(rules[i])
    elif(textIndex [i] == 1):
        author.append(rules[i])
    elif(textIndex [i] == 2):
        affliation.append(rules[i])
    elif(textIndex [i] == 3):
        abstract.append(rules[i])
    elif(textIndex [i] == 4):
        noise.append(rules[i])
    elif(textIndex [i] == 6):
        ids.append(rules[i])
    
        
    
#print(title)





            




print(Filename)
print()
print()
Abstract = {}

for i in range(len(abstract)):
    for j in range(len(abstract[i])):
        value =str(abstract[i][j][0])+","+abstract[i][j][1]+","+str(abstract[i][j][2])
        #print(value)
        
        if(value not in Abstract):
            Abstract[value] = 1
        if(value in Abstract):
            Abstract[value] += 1
        
     


print(sorted(Abstract.items(), key=lambda x: x[1], reverse=True))

#print(Abstract)

'''
for i in Abstract:
    Abstract[i] = Abstract[i] / len(abstract)
    
#print(Abstract)
'''

Title = {}

for i in range(len(title)):
    for j in range(len(title[i])):
        value =str(title[i][j][0])+","+title[i][j][1]+","+str(title[i][j][2])
        #print(value)
        
        if(value not in Title):
            Title[value] = 1
        if(value in Title):
            Title[value] += 1
            
            
print()

print(sorted(Title.items(), key=lambda x: x[1], reverse=True))
#print(Title)
print()
print()
Author = {}

for i in range(len(author)):
    for j in range(len(author[i])):
        value =str(author[i][j][0])+","+author[i][j][1]+","+str(author[i][j][2])
        #print(value)
        
        if(value not in Author):
            Author[value] = 1
        if(value in Author):
            Author[value] += 1

print(sorted(Author.items(), key=lambda x: x[1], reverse=True))
print()
#print(Author)
print()
print()

Affiliation = {}

for i in range(len(affliation)):
    for j in range(len(affliation[i])):
        value =str(affliation[i][j][0])+","+affliation[i][j][1]+","+str(affliation[i][j][2])
        #print(value)
        
        if(value not in Affiliation  ):
            Affiliation [value] = 1
        if(value in Affiliation ):
            Affiliation [value] += 1
            

print(sorted(Affiliation.items(), key=lambda x: x[1], reverse=True))
#print(Affiliation)
print()
print()

Noise = {}

for i in range(len(noise)):
    for j in range(len(noise[i])):
        value =str(noise[i][j][0])+","+noise[i][j][1]+","+str(noise[i][j][2])
        #print(value)
        
        if(value not in  Noise):
            Noise[value] = 1
        if(value in Noise ):
            Noise[value] += 1
            
        
print(sorted(Noise.items(), key=lambda x: x[1], reverse=True))
#print(Noise)



IDS = {}

for i in range(len(ids)):
    for j in range(len(ids[i])):
        value =str(ids[i][j][0])+","+ids[i][j][1]+","+str(ids[i][j][2])
        #print(value)
        
        if(value not in  IDS):
            IDS[value] = 1
        if(value in IDS ):
            IDS[value] += 1
            
#print("IDS Cluster")

#print(IDS)




removeRules = []

maxAbstract =0
AbstractRule = ""
for i in Abstract:
    if(Abstract[i] > maxAbstract):
        maxAbstract =Abstract[i]
        AbstractRule = i
        
        
removeRules.append(AbstractRule)

maxAuthor =0
AuthorRule = ""
for i in Author:
    if(Author[i] > maxAuthor):
        maxAuthor =Author[i]
        AuthorRule = i
        
        
removeRules.append(AuthorRule)


maxAffiliation =0
AffiliationRule = ""
for i in Affiliation:
    if(Affiliation[i] > maxAffiliation):
        maxAffiliation =Affiliation[i]
        AffiliationRule = i
        
        
removeRules.append(AffiliationRule)



















max = 0

ImportantRule=""
for i in Title:
    if(Title[i] > max) and (i not in removeRules)  :
        max = Title[i]
        ImportantRule = i

        
ImportantRuleDiscription = ImportantRule.split(",")
print("IMPORTANT RULE FOR TITLE")
#print(ImportantRuleDiscription)
ImportantRuleDiscription[0] = int(ImportantRuleDiscription[0] )
ImportantRuleDiscription[2] = int(ImportantRuleDiscription[2])



docPreview = docx.Document()


# for i in range(len(text)):
#     for k in range(len(rules[i])):
#         if(rules[i][k] == ImportantRuleDiscription  and len(text[i].split()) >2):
#             print("************************************Title***************************************************************")
            
#     print(text[i])
    
    
# for i in range(len(text)):
#     for k in range(len(rules[i])):
#         if(rules[i][k] == ImportantRuleDiscription  and len(text[i].split()) >2):
#             docPreview.add_paragraph("************************************Title***************************")
            
#     docPreview.add_paragraph(text[i])
    

subHeadingsWords =  ["Introduction:","Objectives","Methods","Results","Conclusion","Background:","Objectives:","Methods:","Results:","Conclusion:","DISCUSSION:"]
    
    
for i in range(len(text)):
    for k in range(len(rules[i])):
        if(rules[i][k] == ImportantRuleDiscription):
            subHeading = text[i].split()
            #print(subHeading)
            if(subHeading[0] not in subHeadingsWords):
                print("************************************Title***************************************************************")
            
    print(text[i])
    
    
for i in range(len(text)):
    for k in range(len(rules[i])):
        if(rules[i][k] == ImportantRuleDiscription):
            subHeading = text[i].split()
            print(subHeading[0])
            if(subHeading[0] not in subHeadingsWords):
                docPreview.add_paragraph("************************************Title***************************")
            
    docPreview.add_paragraph(text[i])
    
    
docPreview.save("C:/Users/darshanRaghunath/"+Filename+"Preview"+".docx")


# indexs = []
# for i in range(len(text)):
#     for k in range(len(rules[i])):
#         if(rules[i][k] == ImportantRuleDiscription  and len(text[i].split()) >2):
#             indexs.append(i)

            
# print(indexs)


indexs = []
for i in range(len(text)):
    for k in range(len(rules[i])):
        if(rules[i][k] == ImportantRuleDiscription):
            subHeading = text[i].split()
            print(subHeading)
            
            '''
            if(subHeading[0] not in subHeadingsWords):
                indexs.append(i)
            '''

            
print(indexs)



Boundary = []
for i in range(len(indexs)-1):
    Box =[]
    Box.append(indexs[i])
    Box.append(indexs[i+1] -1)
    Boundary.append(Box)

Box = [indexs[len(indexs)-1] ,len(text)-1]
Boundary.append(Box)
#print(Boundary)
    

#print(len(text))


I=0
par =0
for page in doc1:
        pageNumber += 1 
        #print("H")
        blocks = page.getText("dict")["blocks"]
        for b in blocks:
            paragraph = ""
            character = []
            if b['type'] == 0:  # block contains text
                #print(b)
                #print()
                #print(par)
                #print(Boundary[I][0])
                if(par == Boundary[I][0]): 
                    #print("H")
                    doc = docx.Document()
                         
                
                if(par>= Boundary[I][0] and    par<= Boundary[I][1]):  
                    doc_para = doc.add_paragraph()
                    for l in b["lines"]:  # iterate through the text lines
                        #print("line")
                        #print(l)
                        for s in l["spans"]:  
                            #print(s)
                            if(int(s["size"]) <=6):
                                super_text = doc_para.add_run(s["text"]+"  ")
                                super_text.font.superscript = True
                                print(s)
                            else:
                                doc_para.add_run("  "+s["text"])
                                
                    
                    
                    
                    
                if(par == Boundary[I][1]):
                    doc.save("C:/Users/darshanRaghunath/"+Filename+"/"+str(I+1)+".docx")
                    I +=1
                    
                par +=1 
       
                    

from docx2pdf import convert

#convert("C:/Users/darshanRaghunath/Doc16new/1.docx")
#convert("C:/Users/darshanRaghunath/Doc14/1.docx")

for i in range(1,2):
    convert("C:/Users/darshanRaghunath/Doc18/"+str(i)+".docx")
    print("C:/Users/darshanRaghunath/Doc18/"+str(i)+".docx")
    







    














